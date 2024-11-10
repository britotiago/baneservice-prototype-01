from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
from django.conf import settings
import os
import json

# Gather necessary data for the document
def gather_data():
    # Define the file path of the existing JSON file
    file_name = 'merged_output.json'
    file_path = os.path.join(os.getcwd(), file_name)

    # Check if the file exists
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"The file '{file_name}' does not exist in the current directory.")

    # Load the JSON data from the existing file
    with open(file_path, 'r', encoding='utf-8') as json_file:
        data = json.load(json_file)

    # Return the loaded JSON data
    return data

def add_header(document):
    section = document.sections[0]
    section.header_distance = Pt(0)
    header = section.header

    # Calculate the available width (page width minus margins)
    available_width = section.page_width - section.left_margin - section.right_margin

    # Create a table in the header that uses the available width
    header_table = header.add_table(rows=1, cols=2, width=available_width)
    header_table.autofit = False

    # Set the table width
    tbl = header_table._tbl
    tblPr = tbl.tblPr
    tblWidth = OxmlElement('w:tblW')
    tblWidth.set(qn('w:w'), str(int(available_width)))
    tblWidth.set(qn('w:type'), 'dxa')
    tblPr.append(tblWidth)

    # Set the column widths
    header_table.columns[0].width = Inches(available_width * 0.8)  # Adjust as needed
    header_table.columns[1].width = Inches(available_width * 0.2)  # Adjust as needed

    header_cells = header_table.rows[0].cells

    # Vertically align both cells
    for cell in header_cells:
        tc_pr = cell._element.get_or_add_tcPr()
        tc_v_align = OxmlElement('w:vAlign')
        tc_v_align.set(qn('w:val'), 'center')
        tc_pr.append(tc_v_align)

    # Left cell with text
    left_paragraph = header_cells[0].add_paragraph()
    run = left_paragraph.add_run("BREEAM Infrastruktur\n")
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.bold = False

    run_italic = left_paragraph.add_run("Samsvarsbeskrivelse")
    run_italic.font.size = Pt(11)
    run_italic.font.name = 'Calibri'
    run_italic.font.italic = True
    left_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Right cell with image
    logo_path = os.path.join(settings.BASE_DIR, 'assets', 'docx_files', 'logo.jpg')
    right_paragraph = header_cells[1].add_paragraph()
    run = right_paragraph.add_run()
    run.add_picture(logo_path, width=Inches(1))
    right_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

def set_cell_font(cell, font_name='Calibri', font_size=12, bold=False):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.bold = bold
        # Apply font for complex scripts like Calibri
        rFonts = paragraph._element.xpath('.//w:rFonts')
        for rFont in rFonts:
            rFont.set(qn('w:eastAsia'), font_name)

def add_project_information_table(document, data):
    section = document.sections[0]
    available_width = section.page_width - section.left_margin - section.right_margin

    project_info_table = document.add_table(rows=4, cols=2)
    project_info_table.style = 'Table Grid'
    project_info_table.autofit = False

    # Set table width
    tbl = project_info_table._tbl
    tblPr = tbl.tblPr
    tblWidth = OxmlElement('w:tblW')
    tblWidth.set(qn('w:w'), str(int(available_width)))
    tblWidth.set(qn('w:type'), 'dxa')
    tblPr.append(tblWidth)

    # Set column widths
    project_info_table.columns[0].width = Inches(available_width * 0.4)
    project_info_table.columns[1].width = Inches(available_width * 0.6)

    project_info_table.cell(0, 0).text = 'Prosjekt:'
    project_info_table.cell(0, 1).text = data['project_name']
    project_info_table.cell(1, 0).text = 'BREEAM Inf. ansvarlig hos entreprenør:'
    project_info_table.cell(1, 1).text = data['breeam_entrepreneur_responsible']
    project_info_table.cell(2, 0).text = 'BREEAM Inf. ansvarlig hos byggherre:'
    project_info_table.cell(2, 1).text = data['breeam_civil_engineer_responsible']
    project_info_table.cell(3, 0).text = 'BREEAM Assessor:'
    project_info_table.cell(3, 1).text = data['breeam_assessor']

    # Apply formatting for Project Information Table
    for row in project_info_table.rows:
        set_cell_font(row.cells[0], bold=True)  # Bold the first column
        set_cell_font(row.cells[1])  # Set the second column without bold

def add_audit_information_table(document, data):
    # Audit Information Table
    audit_info_table = document.add_table(rows=7, cols=2)
    audit_info_table.style = 'Table Grid'
    audit_info_table.cell(0, 0).text = 'Revisjonskriteria:'
    audit_info_table.cell(0, 1).text = data['audit_criteria']
    audit_info_table.cell(1, 0).text = 'Premiss (ja/nei):'
    audit_info_table.cell(1, 1).text = data['premise']
    audit_info_table.cell(2, 0).text = 'Poeng: (n av n totalt)'
    audit_info_table.cell(2, 1).text = data['total_points']
    audit_info_table.cell(3, 0).text = 'Utarbeidet av:'
    audit_info_table.cell(3, 1).text = data['prepared_by']
    audit_info_table.cell(4, 0).text = 'Opprettet dato:'
    audit_info_table.cell(4, 1).text = data['date_created']
    audit_info_table.cell(5, 0).text = 'Dato 1. levering:'
    audit_info_table.cell(5, 1).text = ''
    audit_info_table.cell(6, 0).text = 'Siste revisjonsdato:'
    audit_info_table.cell(6, 1).text = ''

    # Apply formatting for Audit Information Table
    for row in audit_info_table.rows:
        set_cell_font(row.cells[0], bold=True)  # Bold the first column
        set_cell_font(row.cells[1])

# Function to set heading style
def set_heading_style(heading, font_name='Calibri', font_size=11, bold=True, color=(0, 0, 0)):
    run = heading.runs[0]
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(color[0], color[1], color[2])

def add_compliance_description(document, data):
    compliance_heading = document.add_heading('SAMSVARSBESKRIVELSE', level=2)
    compliance_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    set_heading_style(compliance_heading)

    first_entry = True  # Flag to check if it's the first entry

    for entry in data['compliance_description']:
        # Create a paragraph and manually insert a hyphen as the bullet point
        compliance_paragraph = document.add_paragraph()
        # Add a hyphen followed by a tab for separation
        compliance_paragraph.add_run('-\t').font.name = 'Calibri'

        # Add summary text
        summary_run = compliance_paragraph.add_run(entry['summary'])
        summary_run.font.name = 'Calibri'
        summary_run.font.size = Pt(12)

        # Customize the paragraph format for alignment
        paragraph_format = compliance_paragraph.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY  # Set justification
        paragraph_format.left_indent = Pt(40)  # Increase space between the border and the hyphen
        paragraph_format.first_line_indent = Pt(-20)  # Adjust this to fine-tune the space between hyphen and text
        paragraph_format.space_after = Pt(0)  # Set consistent space after paragraph

        if first_entry:
            paragraph_format.space_before = Pt(12)  # Increase space before only the first paragraph
            first_entry = False  # Reset the flag after applying to the first entry
        else:
            paragraph_format.space_before = Pt(0)  # No added space before subsequent paragraphs

        # Adjusting space between hyphen and text
        compliance_paragraph.paragraph_format.tab_stops.add_tab_stop(Pt(40))

def set_cell_font2(cell, font_name='Calibri', bold=True, font_size=12):
    paragraphs = cell.paragraphs
    for paragraph in paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.bold = bold
            run.font.size = Pt(font_size)

def add_attachments_table(document, data):
    # Attachments Table
    attachments_heading = document.add_heading('VEDLEGG', level=2)
    attachments_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    attachments_heading.paragraph_format.space_after = Pt(12)
    set_heading_style(attachments_heading)  # Ensure this function sets the desired styles.

    attachments_table = document.add_table(rows=len(data['attachments']) + 1, cols=3)
    attachments_table.style = 'Table Grid'

    # Header row
    hdr_cells = attachments_table.rows[0].cells
    hdr_cells[0].text = 'Nr.'
    hdr_cells[1].text = 'Navn (husk kobling til dokument)'
    hdr_cells[2].text = 'Beskrivelse'

    # Apply font settings to header using set_cell_font2
    for cell in hdr_cells:
        set_cell_font2(cell)  # Assuming this function sets font to Calibri, bold, and size 12 correctly

    # Attachment rows
    for idx, attachment in enumerate(data['attachments'], 1):
        row_cells = attachments_table.row_cells(idx)
        row_cells[0].text = str(attachment['number'])  # Ensure number is converted to string
        row_cells[1].text = attachment['name']
        row_cells[2].text = attachment['description']

        # Apply font settings to row cells using set_cell_font2
        for cell in row_cells:
            set_cell_font2(cell)  # Use set_cell_font2 to apply font settings

    # Set specific width to the columns
    first_col_width = Pt(40)  # Width for 'Nr.'
    second_col_width = Pt(230)  # Width for 'Navn'
    third_col_width = Pt(230)  # Width for 'Beskrivelse'
    for row in attachments_table.rows:
        row.cells[0].width = first_col_width
        row.cells[1].width = second_col_width
        row.cells[2].width = third_col_width

def adjust_page_margins(document):
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def add_footer(document):
    section = document.sections[0]
    footer = section.footer

    # Add paragraph to footer
    footer_paragraph = footer.paragraphs[0]

    # Create a new run in the paragraph
    run = footer_paragraph.add_run()
    run.font.name = 'Calibri'
    run.font.size = Pt(11)

    # Create a field for page numbering (PAGE field)
    fld_simple = OxmlElement('w:fldSimple')
    fld_simple.set(qn('w:instr'), 'PAGE')

    # Create the run for the field
    run_elem = OxmlElement('w:r')
    run_elem_rpr = OxmlElement('w:rPr')

    # Set the font name and size
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')
    rFonts.set(qn('w:eastAsia'), 'Calibri')
    run_elem_rpr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '22')  # Font size = 11 pt (22 half-points)
    run_elem_rpr.append(sz)

    run_elem.append(run_elem_rpr)
    fld_simple.append(run_elem)

    # Append the field to the footer
    footer_paragraph._element.append(fld_simple)

    # Center the paragraph in the footer
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def create_word_document(data, output_path):
    document = Document()
    adjust_page_margins(document)

    add_header(document)
    add_project_information_table(document, data)
    document.add_paragraph()
    add_audit_information_table(document, data)
    document.add_paragraph()
    add_compliance_description(document, data)
    document.add_paragraph()
    add_attachments_table(document, data)
    add_footer(document)

    # Save the document
    document.save(output_path)

# Main logic
def main():
    data = gather_data()
    output_path = '../../assets/generated_reports/generated_audit_report.docx'  # Updated for local testing
    create_word_document(data, output_path)
    print(f"Document saved at: {output_path}")

if __name__ == "__main__":
    main()